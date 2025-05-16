import streamlit as st
from datetime import date
from streamlit_drawable_canvas import st_canvas
from PIL import Image
st.set_page_config(page_title="Site Sign-Off Checklist", layout="centered")
if "submitted" not in st.session_state:
    st.session_state.submitted = False

if st.session_state.submitted:
    st.success("✅ Form submitted and saved successfully!")
    st.balloons()
    st.session_state.submitted = False  # reset flag
    import time
    time.sleep(2)
    st.experimental_rerun()


st.title("📝 Site Sign-Off Checklist")

# ✅ Define checklist items for each type (put this near the top)
checklists = {
    "Joinery": [
        "Fire Door Installation",
        [
            "Entrance door installed, plumb, and square",
            "Fire rating label visible and undamaged",
            "Intumescent seals fitted correctly (frame or door as required)",
            "Gap tolerances compliant and packers fitted correctly",
            "Door closes fully and latches without resistance",
            "Door furniture/ironmongery fitted to spec"
        ],
        "Door Liners & Pod Door Flat Liners",
        [
            "Entrance door liner installed straight and securely fixed",
            "Pod door flat liner fitted correctly and square"
        ],
        "Internal Doors",
        [
            "Internal doors installed correctly with consistent gaps",
            "Hinges, handles, and latches fitted securely"
        ],
        "Skirting Installation",
        [
            "Skirting fitted tight to wall, level, and flush",
            "Joints (mitres, scribes) tight and filled where required",
            "Fixings fitted neatly"
        ],
        "Architraves",
        [
            "Architraves fitted tight to frame and wall",
            "Consistent margins and clean mitres",
            "No gaps, damage, or poor cuts visible"
        ],
        "Thresholds / Trims",
        [
            "Correct type and size of thresholds installed",
            "Fitted level and securely fixed",
            "No sharp edges, trip hazards, or gaps"
        ],
        "Doorstops",
        ["Doorstops fitted correctly"],
        "Ironmongery",
        [
            "All ironmongery installed to spec",
            "Function tested – doors open, close, and latch smoothly with receivers adjusted",
            "No missing screws or loose components"
        ],
        "Final Checks",
        [
            "Damaged/missing components reported",
            "Area left tidy and safe",
            "Installation ready for inspection/handover"
        ]
    ],
    "Laminate Flooring": [
        "Underlay Installation",
        [
            "Underlay rolled out and cut neatly with no overlapping",
            "Underlay reaches wall edges or per manufacturer’s instruction",
            "No debris trapped underneath"
        ],
        "Laminate Flooring Installation",
        [
            "First row straight and expansion gap allowed (8–10mm or as required)",
            "Correct staggered pattern followed (min 300mm between short joints)",
            "Click-lock or tongue-and-groove joints properly engaged",
            "No visible gaps between boards",
            "Cuts made cleanly and edges tight",
            "Door frames undercut or edge neatly scribed",
            "Expansion gap left around all walls and fixed objects",
            "Thresholds/trims installed as required"
        ],
        "After Installation",
        [
            "Entire floor cleaned and checked for damage",
            "Waste and off-cuts removed from site",
            "Photos taken for records",
            "Snags complete & subcontractor confirms floor complete and ready for sign-off"
        ]
    ]
}

# ✅ Select checklist type
checklist_type = st.selectbox("Select Checklist Type", ["Joinery", "Laminate Flooring"])

# =========================
# 📋 Project Information
# =========================
st.subheader("Project Information")

project_name = st.text_input("Project Name / Block", key="project_name")
unit_number = st.text_input("Apartment / Unit Number", key="unit_number")
inspection_date = st.date_input("Date of Inspection", value=date.today(), key="inspection_date")

# =========================
# 👷 Personnel Information
# =========================
st.subheader("Personnel")

subcontractor_name = st.text_input("Subcontractor Name", key="subcontractor_name")
foreman_name = st.text_input("CH Foreman", key="foreman_name")

# ✅ Preserve form state
if "form_filled" not in st.session_state:
    st.session_state.form_filled = False

# ✅ Button to unlock checklist
if st.button("Next Step"):
    st.session_state.form_filled = True

# ✅ Show checklist and signature canvases after form submission
if st.session_state.form_filled:
    st.success("Form filled successfully. Now complete the checklist below.")

    checklist_data = []
    for item in checklists[checklist_type]:
        if isinstance(item, str):
            st.subheader(item)
        elif isinstance(item, list):
            for sub_item in item:
                checked = st.checkbox(sub_item, key=sub_item)
                checklist_data.append((sub_item, checked))



    st.caption("✍️ Please sign below — must not be left blank.")

    st.subheader("📌 Subcontractor Signature")
    canvas_result = st_canvas(
        fill_color="rgba(255, 255, 255, 0)",  # Transparent background
        stroke_width=2,
        stroke_color="#000000",
        background_color="#ffffff",
        height=150,
        width=400,
        drawing_mode="freedraw",
        key="signature",
    )

    # Capture the subcontractor signature if drawn
    signature_img_path = None
    if canvas_result.image_data is not None and canvas_result.image_data.any():
        signature_img = Image.fromarray(canvas_result.image_data.astype("uint8"))
        signature_img_path = "signature.png"
        signature_img.save(signature_img_path)

    st.subheader("📌 CH Foreman Signature")
    canvas_result_foreman = st_canvas(
        fill_color="rgba(255, 255, 255, 0)",  # Transparent background
        stroke_width=2,
        stroke_color="#000000",
        background_color="#ffffff",
        height=150,
        width=400,
        drawing_mode="freedraw",
        key="foreman_signature",
    )

    # Capture the foreman signature if drawn
    foreman_signature_img_path = None
    if canvas_result_foreman.image_data is not None and canvas_result_foreman.image_data.any():
        foreman_signature_img = Image.fromarray(canvas_result_foreman.image_data.astype("uint8"))
        foreman_signature_img_path = "foreman_signature.png"
        foreman_signature_img.save(foreman_signature_img_path)


import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# Save path
SAVE_DIR = "C:/Users/abinash.parichha/signoff_checklist"

if st.button("Generate DOCX and Save"):
    doc = Document()

    # ========== HEADER ========== #
    logo_path = "logo.png"  # Your actual logo path
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(2.5))
        doc.add_paragraph()  # spacing after logo

    # Add centered title
    title = doc.add_heading("Joinery Installation Sign-Off Sheet", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # extra space below title

    # ========== PROJECT INFO ========== #
    info = doc.add_paragraph()
    info.style.font.size = Pt(11)
    info.add_run("Checklist Type: ").bold = True
    info.add_run(f"{checklist_type}\n")
    info.add_run("Project Name / Block: ").bold = True
    info.add_run(f"{project_name}\n")
    info.add_run("Apartment / Unit: ").bold = True
    info.add_run(f"{unit_number}\n")
    info.add_run("Date of Inspection: ").bold = True
    info.add_run(f"{inspection_date}\n")
    info.add_run("Subcontractor Name: ").bold = True
    info.add_run(f"{subcontractor_name}\n")
    info.add_run("CH Foreman: ").bold = True
    info.add_run(f"{foreman_name}\n")

    doc.add_paragraph()

       # ========== CHECKLIST ========== #
    doc.add_heading("Checklist Items", level=1)

    # Get the list of section titles
    section_titles = [sec for sec in checklists[checklist_type] if isinstance(sec, str)]

    section_index = -1
    for item in checklists[checklist_type]:
        if isinstance(item, str):
            # It's a section title
            doc.add_paragraph()  # spacing
            doc.add_heading(item, level=3)
            section_index += 1
        elif isinstance(item, list):
            for question in item:
                # Find matching checked status from checklist_data
                matched = next((val for text, val in checklist_data if text == question), False)
                symbol = "✅" if matched else "☐"
                line = doc.add_paragraph(style='List Bullet')
                run = line.add_run(f"{question} {symbol}")
                run.font.size = Pt(10)



    # ========== SIGNATURES ========== #
    doc.add_paragraph()
    doc.add_paragraph("Subcontractor Signature:")
    if signature_img_path and os.path.exists(signature_img_path):
        doc.add_picture(signature_img_path, width=Inches(2))

    doc.add_paragraph()
    doc.add_paragraph("CH Foreman Signature:")
    if foreman_signature_img_path and os.path.exists(foreman_signature_img_path):
        doc.add_picture(foreman_signature_img_path, width=Inches(2))

    # ========== SAVE .DOCX (No PDF conversion) ========== #
    filename_base = f"{checklist_type}_{project_name}_{unit_number}_{inspection_date}"
    docx_file = f"{filename_base}.docx"
    doc.save(docx_file)

    # Move DOCX to save folder
    final_path = os.path.join(SAVE_DIR, docx_file)
    os.replace(docx_file, final_path)

    # Clean up
    if signature_img_path and os.path.exists(signature_img_path):
        os.remove(signature_img_path)
    if foreman_signature_img_path and os.path.exists(foreman_signature_img_path):
        os.remove(foreman_signature_img_path)

    st.success(f"✅ Checklist saved as Word file at:\n{final_path}")
    st.session_state.submitted = True

